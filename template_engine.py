"""
Движок обработки шаблонов .docx.

═══════════════════════════════════════════════════════════════
СИНТАКСИС ШАБЛОНА
═══════════════════════════════════════════════════════════════

1. ТЕКСТОВЫЕ ПОЛЯ (для текущего студента):
   {{ ФИО }}              — значение поля «ФИО»
   {{ Группа }}           — значение поля «Группа»

2. ТАБЛИЦА ВСЕХ СТУДЕНТОВ (через цикл в строке таблицы):
   В строке таблицы разместите:
   {% for s in students %}{{ s.ФИО }} {{ s.Группа }}{% endfor %}

   Строка будет продублирована для каждого студента.
"""

import copy
import re
import shutil
import tempfile
from pathlib import Path
from typing import Optional

from docx import Document
from docxtpl import DocxTemplate

import config
from database import Database

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Паттерн для поиска Jinja2-переменных: {{ имя }}
_VAR_PATTERN = re.compile(r"\{\{\s*([^%{}\s][^{}]*?)\s*\}\}")


def find_placeholders(docx_path: Path) -> list[str]:
    """Находит все плейсхолдеры {{ ... }} в документе."""
    doc = Document(str(docx_path))
    placeholders: set[str] = set()

    def _scan(text: str) -> None:
        for m in _VAR_PATTERN.finditer(text):
            name = m.group(1).strip()
            if name.startswith(("%", "tr ", "for ", "if ", "end")):
                continue
            placeholders.add(name)

    for para in doc.paragraphs:
        _scan(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _scan(para.text)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                for para in hf.paragraphs:
                    _scan(para.text)

    return sorted(placeholders)


def _build_student_dict(
    row_data: dict[str, str],
    original_headers_map: Optional[dict[str, str]],
) -> dict[str, str]:
    """Строит словарь полей студента с оригинальными и санированными именами."""
    student: dict[str, str] = {}
    if original_headers_map:
        for orig, safe in original_headers_map.items():
            value = str(row_data.get(safe, ""))
            student[orig] = value
            student[safe] = value
    else:
        student = {k: str(v) for k, v in row_data.items()}
    return student


def _flatten_cell_to_single_paragraph(cell_elem) -> None:
    """
    Объединяет все параграфы ячейки в один параграф с одним текстовым run.
    Используется для строк с циклами, чтобы regex корректно работал с Jinja-тегами,
    разбитыми переносами строк внутри ячейки.
    """
    paragraphs = cell_elem.findall(f'{W_NS}p')
    if not paragraphs:
        return

    # Собираем весь текст из всех параграфов, разделяя их переводом строки
    parts = []
    for p in paragraphs:
        texts = p.findall(f'.//{W_NS}t')
        para_text = ''.join(t.text or '' for t in texts)
        parts.append(para_text)
    full_text = '\n'.join(parts)

    # Оставляем первый параграф, объединяя все run'ы в один
    first_p = paragraphs[0]
    texts = first_p.findall(f'.//{W_NS}t')
    if texts:
        texts[0].text = full_text
        texts[0].set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        for t in texts[1:]:
            t.text = ''
    else:
        # Нет существующих <w:t> — создаём новый run
        from docx.oxml import OxmlElement
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = full_text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        first_p.append(r)

    # Удаляем остальные параграфы
    for p in paragraphs[1:]:
        cell_elem.remove(p)


def _find_loop_row(table) -> Optional[int]:
    """Находит индекс строки, содержащей {% for s in students %}."""
    for row_idx, row in enumerate(table.rows):
        row_text = '\n'.join(cell.text for cell in row.cells)
        if re.search(r'\{%\s*(?:tr\s+)?for\s+\w+\s+in\s+\w+', row_text):
            return row_idx
    return None


def _get_cell_grid_span(tc) -> int:
    """Возвращает gridSpan ячейки (сколько столбцов она занимает)."""
    tcPr = tc.find(f'{W_NS}tcPr')
    if tcPr is not None:
        gridSpan = tcPr.find(f'{W_NS}gridSpan')
        if gridSpan is not None:
            val = gridSpan.get(f'{W_NS}val') or gridSpan.get('val')
            if val:
                try:
                    return int(val)
                except ValueError:
                    pass
    return 1


def _set_cell_grid_span(tc, span: int) -> None:
    """Устанавливает gridSpan ячейки."""
    from docx.oxml import OxmlElement
    tcPr = tc.find(f'{W_NS}tcPr')
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    gridSpan = tcPr.find(f'{W_NS}gridSpan')
    if gridSpan is None:
        gridSpan = OxmlElement('w:gridSpan')
        tcPr.append(gridSpan)
    gridSpan.set(f'{W_NS}val', str(span))


def _cell_is_empty(tc) -> bool:
    """Проверяет, пуста ли ячейка (нет текста после подстановки)."""
    for t in tc.findall(f'.//{W_NS}t'):
        if t.text and t.text.strip():
            return False
    return True


def _merge_empty_cells_left(row_element) -> None:
    """
    Объединяет пустые ячейки с левой соседней: если ячейка пуста, её
    gridSpan добавляется к левой непустой ячейке, а сама пустая удаляется.
    """
    tcs = row_element.findall(f'{W_NS}tc')
    # Идём справа налево, чтобы безопасно удалять элементы
    i = len(tcs) - 1
    while i >= 1:
        tc = tcs[i]
        if _cell_is_empty(tc):
            left_tc = tcs[i - 1]
            left_span = _get_cell_grid_span(left_tc)
            this_span = _get_cell_grid_span(tc)
            _set_cell_grid_span(left_tc, left_span + this_span)
            row_element.remove(tc)
        i -= 1


def _merge_empty_cells_right(row_element) -> None:
    """
    Объединяет пустые ячейки с правой соседней: если ячейка пуста, её
    gridSpan добавляется к правой ячейке, а сама пустая удаляется.
    """
    tcs = row_element.findall(f'{W_NS}tc')
    i = 0
    while i < len(tcs) - 1:
        tc = tcs[i]
        if _cell_is_empty(tc):
            right_tc = tcs[i + 1]
            right_span = _get_cell_grid_span(right_tc)
            this_span = _get_cell_grid_span(tc)
            _set_cell_grid_span(right_tc, right_span + this_span)
            row_element.remove(tc)
            tcs.pop(i)
            # не инкрементируем i — следующий элемент теперь на той же позиции
        else:
            i += 1


def _make_group_header_row(template_row_elem, group_label: str, total_cols: int):
    """
    Создаёт строку-заголовок группы: копия шаблонной строки с одной объединённой
    ячейкой, содержащей текст group_label.
    """
    row_copy = copy.deepcopy(template_row_elem)
    tcs = row_copy.findall(f'{W_NS}tc')
    if not tcs:
        return row_copy

    # Оставляем первую ячейку, удаляем остальные
    first_tc = tcs[0]
    for tc in tcs[1:]:
        row_copy.remove(tc)

    # Расширяем первую ячейку на все столбцы
    _set_cell_grid_span(first_tc, total_cols)

    # Заменяем содержимое первой ячейки на group_label
    _flatten_cell_to_single_paragraph(first_tc)
    for t in first_tc.findall(f'.//{W_NS}t'):
        t.text = group_label
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        break
    else:
        # Нет текстового элемента — создаём
        from docx.oxml import OxmlElement
        p_new = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = group_label
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        p_new.append(r)
        first_tc.append(p_new)

    # Выравниваем текст заголовка по левому краю
    from docx.oxml import OxmlElement as _OxmlElement
    p = first_tc.find(f'{W_NS}p')
    if p is not None:
        pPr = p.find(f'{W_NS}pPr')
        if pPr is None:
            pPr = _OxmlElement('w:pPr')
            p.insert(0, pPr)
        jc = pPr.find(f'{W_NS}jc')
        if jc is None:
            jc = _OxmlElement('w:jc')
            pPr.append(jc)
        jc.set(f'{W_NS}val', 'left')

    return row_copy


def _get_table_grid_cols(table_elem) -> int:
    """Возвращает количество столбцов таблицы по <w:tblGrid>."""
    tblGrid = table_elem.find(f'{W_NS}tblGrid')
    if tblGrid is not None:
        return len(tblGrid.findall(f'{W_NS}gridCol'))
    # Fallback — суммарный gridSpan первой строки
    first_tr = table_elem.find(f'{W_NS}tr')
    if first_tr is not None:
        return sum(_get_cell_grid_span(tc) for tc in first_tr.findall(f'{W_NS}tc'))
    return 1


def _get_student_field(student: dict, field_name: str) -> str:
    """Возвращает значение поля студента, пробуя оригинальное и санированное имя."""
    val = student.get(field_name, '') or ''
    if not val:
        sanitized = re.sub(r'\s+', '_', field_name.strip()).lower()
        val = student.get(sanitized, '') or ''
    return val.strip() or '(без значения)'


def _group_students_recursive(
    students_list: list[dict],
    group_by_list: list[str],
) -> list:
    """
    Рекурсивно группирует студентов по нескольким уровням.

    Возвращает:
    - list[dict] если group_by_list пуст (студенты напрямую)
    - list[tuple[str, list]] если group_by_list не пуст (группы с детьми)
    """
    if not group_by_list:
        return students_list

    col = group_by_list[0]
    rest = group_by_list[1:]

    groups: dict[str, list[dict]] = {}
    order: list[str] = []
    for s in students_list:
        value = _get_student_field(s, col)
        if value not in groups:
            groups[value] = []
            order.append(value)
        groups[value].append(s)

    return [
        (key, _group_students_recursive(groups[key], rest))
        for key in order
    ]


def _expand_grouped_rows(
    parent,
    pos: int,
    tr_element,
    items: list,
    level: int,
    label_templates: list[str],
    total_cols: int,
    merge_empty_cells: bool,
    merge_right: bool,
) -> int:
    """
    Рекурсивно вставляет строки: заголовки групп и строки студентов.
    Возвращает обновлённую позицию вставки.
    """
    def _get_label_template(lvl: int) -> str:
        if label_templates:
            return label_templates[min(lvl, len(label_templates) - 1)]
        return "{value}"

    for item in items:
        if isinstance(item, tuple):
            group_value, children = item
            label = _get_label_template(level).format(value=group_value)
            header = _make_group_header_row(tr_element, label, total_cols)
            parent.insert(pos, header)
            pos += 1
            pos = _expand_grouped_rows(
                parent, pos, tr_element, children, level + 1,
                label_templates, total_cols, merge_empty_cells, merge_right,
            )
        else:
            # item — словарь студента
            tr_copy = copy.deepcopy(tr_element)
            _substitute_student_in_row(tr_copy, item)
            if merge_empty_cells:
                if merge_right:
                    _merge_empty_cells_right(tr_copy)
                else:
                    _merge_empty_cells_left(tr_copy)
            parent.insert(pos, tr_copy)
            pos += 1
    return pos


def _substitute_student_in_row(row_element, student: dict[str, str]) -> None:
    """
    Подставляет значения студента в строку таблицы:
    - {{ s.FieldName }} → фактическое значение
    - Удаляет {% for %} и {% endfor %} markers

    Сначала сплющивает содержимое каждой ячейки в один параграф, чтобы
    regex корректно работал с Jinja-тегами, разбитыми на несколько параграфов.
    """
    def _replace(match, student_dict):
        field_name = match.group(1)
        # Убираем ведущие/конечные пробелы и переносы (но сохраняем внутренние)
        field_name = field_name.strip()
        if field_name in student_dict:
            return student_dict[field_name]
        # Пробуем санированную версию
        sanitized = re.sub(r'\s+', '_', field_name).lower()
        if sanitized in student_dict:
            return student_dict[sanitized]
        return ''

    # Обрабатываем каждую ячейку
    for tc in row_element.findall(f'.//{W_NS}tc'):
        # Сплющиваем ячейку в один параграф
        _flatten_cell_to_single_paragraph(tc)

        # Теперь в ячейке один параграф с одним текстовым run — можем делать regex
        for t in tc.findall(f'.//{W_NS}t'):
            if not t.text:
                continue
            text = t.text

            # Удаляем loop markers (с учётом возможных переносов)
            text = re.sub(r'\{%\s*(?:tr\s+)?for[^%]*?%\}', '', text, flags=re.DOTALL)
            text = re.sub(r'\{%\s*(?:tr\s+)?endfor[^%]*?%\}', '', text, flags=re.DOTALL)

            # Заменяем {{ s.FieldName }} на фактическое значение
            # DOTALL чтобы . матчил переносы (для полей с multi-line именами)
            text = re.sub(
                r'\{\{\s*s\.(.*?)\}\}',
                lambda m: _replace(m, student),
                text,
                flags=re.DOTALL,
            )

            t.text = text


def _expand_table_loops(
    template_path: Path,
    students_list: list[dict],
    group_by: Optional[list[str]] = None,
    group_label_templates: Optional[list[str]] = None,
    merge_empty_cells: bool = True,
    merge_right: bool = False,
) -> None:
    """
    Разворачивает циклы в таблицах:
    - Находит строку с {% for s in students %}
    - Рекурсивно группирует студентов по нескольким уровням (если задан group_by)
    - Для каждой группы вставляет объединённую строку-заголовок
    - Дублирует строку шаблона для каждого студента, подставляя значения
    - Объединяет пустые ячейки (если merge_empty_cells)
    """
    doc = Document(str(template_path))

    group_by_list = group_by or []
    label_tpls = group_label_templates or ["{value}"]

    for table in doc.tables:
        loop_row_idx = _find_loop_row(table)
        if loop_row_idx is None:
            continue

        template_row = table.rows[loop_row_idx]
        tr_element = template_row._element
        parent = tr_element.getparent()
        insert_pos = parent.index(tr_element)

        total_cols = _get_table_grid_cols(table._tbl)

        grouped = _group_students_recursive(students_list, group_by_list)

        _expand_grouped_rows(
            parent, insert_pos, tr_element, grouped, 0,
            label_tpls, total_cols, merge_empty_cells, merge_right,
        )

        parent.remove(tr_element)

    doc.save(str(template_path))


def generate_document(
    template_path: Path,
    row_data: dict[str, str],
    output_filename: str,
    all_rows_data: Optional[tuple[list[str], list[list]]] = None,
    original_headers_map: Optional[dict[str, str]] = None,
    group_by: Optional[list[str]] = None,
    group_label_templates: Optional[list[str]] = None,
    merge_empty_cells: bool = True,
    merge_right: bool = False,
) -> Path:
    """
    Генерирует один документ из шаблона.

    group_by: список колонок для многоуровневой группировки (например, ["Факультет", "Группа"]).
    group_label_templates: шаблон заголовка для каждого уровня, {value} — значение поля.
    merge_empty_cells: объединять пустые ячейки с соседней.
    merge_right: True — объединять с правой, False — с левой.
    """
    # Текущий студент (для прямых плейсхолдеров вне таблиц)
    current_student = _build_student_dict(row_data, original_headers_map)

    # Все студенты (для таблиц с циклами)
    students_list: list[dict[str, str]] = []
    if all_rows_data:
        headers_raw, rows_raw = all_rows_data
        rev = {}
        if original_headers_map:
            rev = {safe: orig for orig, safe in original_headers_map.items()}

        for row in rows_raw:
            s: dict[str, str] = {}
            for h_raw, val in zip(headers_raw, row):
                val_str = str(val) if val is not None else ""
                orig = rev.get(h_raw, h_raw)
                s[orig] = val_str
                s[h_raw] = val_str
            students_list.append(s)

    # Работаем с копией шаблона
    with tempfile.TemporaryDirectory() as tmpdir:
        work_template = Path(tmpdir) / template_path.name
        shutil.copy2(template_path, work_template)

        # Шаг 1: разворачиваем циклы в таблицах (с группировкой и подстановкой)
        _expand_table_loops(
            work_template,
            students_list,
            group_by=group_by,
            group_label_templates=group_label_templates,
            merge_empty_cells=merge_empty_cells,
            merge_right=merge_right,
        )

        # Шаг 2: рендерим оставшиеся плейсхолдеры через docxtpl (нетабличные)
        context: dict = {**current_student}

        tpl = DocxTemplate(str(work_template))
        tpl.render(context, autoescape=False)

        output_path = config.OUTPUT_DIR / output_filename
        tpl.save(str(output_path))

    return output_path


def generate_documents_for_all_rows(
    template_path: Path,
    table_name: str,
    filename_column: Optional[str] = None,
    db: Optional[Database] = None,
    group_by: Optional[list[str]] = None,
    group_label_templates: Optional[list[str]] = None,
    merge_empty_cells: bool = True,
    merge_right: bool = False,
) -> list[Path]:
    """
    Генерирует по одному документу на каждую строку данных в таблице.

    group_by: список колонок для многоуровневой группировки.
    group_label_templates: шаблоны заголовков по одному на каждый уровень.
    """
    close_db = db is None
    if db is None:
        db = Database()

    try:
        headers, all_rows = db.get_all_data(table_name)
        original_map = db.get_original_headers(table_name)
        all_rows_data = (headers, all_rows)

        generated: list[Path] = []

        for idx, row in enumerate(all_rows):
            row_data = dict(zip(headers, row))

            if filename_column and filename_column in row_data:
                raw_name = str(row_data[filename_column])
                safe_name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", raw_name).strip()
                output_name = f"{safe_name or f'document_{idx + 1:03d}'}.docx"
            else:
                output_name = f"document_{idx + 1:03d}.docx"

            path = generate_document(
                template_path=template_path,
                row_data=row_data,
                output_filename=output_name,
                all_rows_data=all_rows_data,
                original_headers_map=original_map,
                group_by=group_by,
                group_label_templates=group_label_templates,
                merge_empty_cells=merge_empty_cells,
                merge_right=merge_right,
            )
            generated.append(path)

        return generated

    finally:
        if close_db:
            db.close()
